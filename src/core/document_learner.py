"""
Document learning module for analyzing existing Gutachten
"""

import re
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
import json

from ..utils.logger import get_logger


class DocumentLearner:
    """
    Learns templates and patterns from existing Gutachten documents
    """
    
    def __init__(self, config):
        """Initialize document learner"""
        self.config = config
        self.logger = get_logger("DocumentLearner")
        
        # Common Gutachten sections in German
        self.sections = [
            "Einleitung", "Vorbemerkung", "Auftrag", "Sachverhalt",
            "Befund", "Beurteilung", "Zusammenfassung", "Gutachten"
        ]
        
        # Medical terminology patterns
        self.medical_patterns = [
            r"Diagnose[:\s]+([^.\n]+)",
            r"Befund[:\s]+([^.\n]+)",
            r"Anamnese[:\s]+([^.\n]+)",
            r"Untersuchung[:\s]+([^.\n]+)"
        ]
    
    def learn_from_document(self, doc_path: str) -> Dict[str, Any]:
        """
        Learn template from a single document
        
        Args:
            doc_path: Path to .doc/.docx file
            
        Returns:
            Dictionary with extracted template data
        """
        self.logger.info(f"Learning from document: {doc_path}")
        
        try:
            # Load document
            doc = Document(doc_path)
            
            # Extract text and structure
            text_content = self._extract_text(doc)
            structure = self._analyze_structure(text_content)
            patterns = self._extract_patterns(text_content)
            formatting = self._analyze_formatting(doc)
            
            template_data = {
                'source_file': doc_path,
                'structure': structure,
                'patterns': patterns,
                'formatting': formatting,
                'text_sample': text_content[:1000],  # First 1000 chars as sample
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'sections_found': len(structure['sections']),
                    'patterns_found': len(patterns)
                }
            }
            
            self.logger.info(f"Successfully learned template from {doc_path}")
            return template_data
            
        except Exception as e:
            self.logger.error(f"Error learning from {doc_path}: {e}")
            raise
    
    def _extract_text(self, doc: Document) -> str:
        """Extract text content from document"""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        return "\n".join(text_parts)
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure and sections"""
        structure = {
            'sections': {},
            'section_order': [],
            'total_length': len(text)
        }
        
        lines = text.split('\n')
        current_section = "Einleitung"
        
        for line in lines:
            line_lower = line.lower()
            
            # Check if line contains section headers
            for section in self.sections:
                if section.lower() in line_lower and len(line.strip()) < 100:
                    current_section = section
                    if section not in structure['section_order']:
                        structure['section_order'].append(section)
                    break
            
            # Add content to current section
            if current_section not in structure['sections']:
                structure['sections'][current_section] = []
            
            if line.strip():
                structure['sections'][current_section].append(line.strip())
        
        return structure
    
    def _extract_patterns(self, text: str) -> List[Dict[str, Any]]:
        """Extract common patterns from text"""
        patterns = []
        
        # Extract medical terminology
        for pattern in self.medical_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                patterns.append({
                    'type': 'medical_terminology',
                    'pattern': pattern,
                    'matches': matches[:5]  # Limit to first 5 matches
                })
        
        # Extract sentence patterns
        sentences = re.split(r'[.!?]+', text)
        sentence_patterns = self._analyze_sentence_patterns(sentences)
        patterns.extend(sentence_patterns)
        
        return patterns
    
    def _analyze_sentence_patterns(self, sentences: List[str]) -> List[Dict[str, Any]]:
        """Analyze sentence structure patterns"""
        patterns = []
        
        # Common German sentence starters
        starters = [
            "Der Patient", "Die Patientin", "Es wurde", "Es zeigt sich",
            "Die Untersuchung", "Der Befund", "Die Diagnose"
        ]
        
        for starter in starters:
            matching_sentences = [s for s in sentences if s.strip().startswith(starter)]
            if matching_sentences:
                patterns.append({
                    'type': 'sentence_starter',
                    'pattern': starter,
                    'count': len(matching_sentences),
                    'examples': matching_sentences[:3]
                })
        
        return patterns
    
    def _analyze_formatting(self, doc: Document) -> Dict[str, Any]:
        """Analyze document formatting"""
        formatting = {
            'paragraphs': [],
            'styles': set(),
            'font_sizes': set(),
            'bold_sections': []
        }
        
        for paragraph in doc.paragraphs:
            para_info = {
                'text': paragraph.text[:100],  # First 100 chars
                'style': paragraph.style.name if paragraph.style else 'Normal',
                'alignment': paragraph.alignment
            }
            formatting['paragraphs'].append(para_info)
            formatting['styles'].add(para_info['style'])
            
            # Check for bold text (section headers)
            for run in paragraph.runs:
                if run.bold and len(run.text.strip()) < 100:
                    formatting['bold_sections'].append(run.text.strip())
        
        # Convert sets to lists for JSON serialization
        formatting['styles'] = list(formatting['styles'])
        formatting['font_sizes'] = list(formatting['font_sizes'])
        
        return formatting
    
    def compare_templates(self, template1: Dict[str, Any], template2: Dict[str, Any]) -> float:
        """
        Compare two templates and return similarity score
        
        Args:
            template1: First template
            template2: Second template
            
        Returns:
            Similarity score (0-1)
        """
        # Simple similarity based on structure and patterns
        structure_similarity = self._compare_structures(
            template1.get('structure', {}),
            template2.get('structure', {})
        )
        
        pattern_similarity = self._compare_patterns(
            template1.get('patterns', []),
            template2.get('patterns', [])
        )
        
        return (structure_similarity + pattern_similarity) / 2
    
    def _compare_structures(self, struct1: Dict, struct2: Dict) -> float:
        """Compare document structures"""
        sections1 = set(struct1.get('section_order', []))
        sections2 = set(struct2.get('section_order', []))
        
        if not sections1 and not sections2:
            return 0.0
        
        intersection = len(sections1.intersection(sections2))
        union = len(sections1.union(sections2))
        
        return intersection / union if union > 0 else 0.0
    
    def _compare_patterns(self, patterns1: List, patterns2: List) -> float:
        """Compare pattern lists"""
        if not patterns1 and not patterns2:
            return 0.0
        
        pattern_types1 = {p.get('type', '') for p in patterns1}
        pattern_types2 = {p.get('type', '') for p in patterns2}
        
        intersection = len(pattern_types1.intersection(pattern_types2))
        union = len(pattern_types1.union(pattern_types2))
        
        return intersection / union if union > 0 else 0.0 